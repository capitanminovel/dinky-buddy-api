"""
Appends new strain profiles to the existing Strain_Profiles.docx.
Existing content is never touched. New strains are grouped by:
  FLOWER / PRE-ROLL / VAPES
and appended at the end of the document.

Usage:
  python generate_strain_doc.py

Output: docs/Strain_Profiles_Updated.docx
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC   = Path("/root/.claude/uploads/0ad01f0e-7b63-4bf4-a73b-2eb47f32e4da/6e75def0-Strain_Profiles.docx")
OUT   = Path("docs/Strain_Profiles_Updated.docx")

# ── Strains already in the document — skip these ─────────────────────────────
ALREADY_IN_DOC = {
    "Cap Junky", "Glue 31", "Egg Roll",
    "Super Lemon Haze", "Alaskan Thunder Fuck", "OG Kush",
    "Green Crack", "Guava", "SuperJack", "Forbidden Fruit",
    "Insane Punch", "Barry White",
}

# ── New strain data — effects based on COA terpene research, not scraped ─────
# Format: key = normalized name, value = dict matching doc field order
STRAINS = {

    # ── FLOWER ──────────────────────────────────────────────────────────────
    "Fight Club": dict(
        category="flower",
        strain_type="Hybrid (Sativa)",
        lineage="The Soap x Velvet Runtz (Campfire Cannabis in-house cross)",
        effects="Euphoric, Creative, Focused, Uplifting, Happy, Cerebral, Relaxing",
        flavors="Fruity, Berry, Citrus, Diesel, Earthy, Pine, Sweet, Woody",
        terpenes="Myrcene, Pinene, Limonene, Caryophyllene, Ocimene, Terpinolene",
        therapeutic="Stress, Fatigue, Depression, Mild Pain, Creative Blocks, Mood Swings",
        negative="Dry Mouth & Eyes, Occasional Anxiety at High Doses",
        aroma="Clean soapy citrus, sweet mixed berry, light diesel backbone, faint pine and earthy woods",
        misc="Campfire Cannabis in-house cross. THC 20–25%. Versatile hybrid — sativa-leaning first hour then settles into relaxed body state. Great for daytime creative work or afternoon social sessions.",
    ),
    "LIT OG": dict(
        category="flower",
        strain_type="Hybrid (Indica)",
        lineage="Tahoe OG x Runtz",
        effects="Relaxing, Sleepy, Euphoric, Calming, Body High, Hungry",
        flavors="Citrus, Creamy, Fruity, Pine, Sweet, Tropical, Woody",
        terpenes="Linalool, Ocimene, Terpinolene, Geraniol, Bisabolol",
        therapeutic="Insomnia, Chronic Pain, Stress, Anxiety, Muscle Tension",
        negative="Dry Mouth & Eyes, Heavy Sedation — Not for Daytime Use",
        aroma="Earthy OG pine with sweet candy citrus, creamy tropical undertones, classic kush musk base",
        misc="Combines Tahoe OG's heavy indica body stone with Runtz's candy sweetness. THC 25–30%. Dense frosty buds. Best for evening wind-down or pre-sleep.",
    ),
    "MAC Stomper": dict(
        category="flower",
        strain_type="Hybrid",
        lineage="MAC (Miracle Alien Cookies) x Grape Stomper (Gage Green Seeds)",
        effects="Euphoric, Relaxing, Sleepy, Giggly, Uplifting, Aroused, Body High",
        flavors="Fruity, Berry, Citrus, Apple, Sweet, Sour, Grape",
        terpenes="Myrcene, Linalool, Pinene, Caryophyllene",
        therapeutic="Chronic Pain, Stress, Insomnia, Mood Disorders, Nausea",
        negative="Dry Mouth & Eyes, Couch-Lock at High Doses",
        aroma="Creamy grape candy, sweet apple, citrus zest, light floral, earthy funk on the finish",
        misc="Cross of two award-winning genetics. MAC contributes creamy citrus and potency; Grape Stomper adds grape and fruit sweetness. THC 24–29%. Available in half-ounce. Dense, trichome-heavy flowers with deep purple and green hues.",
    ),
    "Soap": dict(
        category="flower",
        strain_type="Hybrid (Sativa)",
        lineage="Animal Mints x Kush Mints (Seed Junky Genetics)",
        effects="Euphoric, Focused, Creative, Uplifting, Energetic, Tingly, Aroused",
        flavors="Cheese, Citrus, Mint, Pine, Rose, Sour, Sweet, Tea",
        terpenes="Linalool, Limonene, Caryophyllene, Humulene",
        therapeutic="Depression, Fatigue, Stress, Anxiety, Mild Chronic Pain",
        negative="Dry Mouth & Eyes, Possible Anxiety at High Doses in Novice Users",
        aroma="Clean floral soap, fresh mint, citrus zest, light cheese, pine — one of the most distinctive and refreshing terpene profiles in cannabis",
        misc="Created by Seed Junky Genetics. Famously described as smelling like fresh soap or floral laundry detergent. THC 20–26%. Euphoric and uplifting with a clean, lasting head high. Strong daytime strain for focus and creativity.",
    ),
    "Purple Octane": dict(
        category="flower",
        strain_type="Hybrid (Indica)",
        lineage="(Biscotti x Sherb BX1) x Jealousy F2 (Compound Genetics / Seed Junky)",
        effects="Relaxing, Sleepy, Euphoric, Body High, Tingly, Unbothered",
        flavors="Fruity, Berry, Chemical, Diesel, Sweet, Grape",
        terpenes="Caryophyllene",
        therapeutic="Insomnia, Chronic Pain, Stress, Appetite Stimulation, Muscle Tension",
        negative="Dry Mouth & Eyes, Heavy Sedation — Not for Daytime Use",
        aroma="Sweet grape diesel, rich berry, creamy dessert notes, faint chemical fuel on exhale",
        misc="THC 25–30%+. Dense purple-tinted buds with heavy resin coverage. Hard-hitting indica-dominant — reserved for evening or experienced users.",
    ),
    "Runtz": dict(
        category="flower",
        strain_type="Hybrid",
        lineage="Zkittlez x Gelato (Cookies Fam / Runtz Brand)",
        effects="Euphoric, Happy, Giggly, Uplifting, Social, Relaxing, Sleepy, Hungry, Talkative",
        flavors="Fruity, Citrus, Spice, Sweet, Tropical, Pineapple, Apricot",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Stress, Anxiety, Depression, Chronic Pain, Appetite Stimulation",
        negative="Dry Mouth & Eyes, Possible Sedation at High Doses",
        aroma="Tropical candy sweetness, pineapple citrus, fruity Zkittlez notes, creamy Gelato undertone, light spice on exhale",
        misc="One of the most influential strains of the 2020s. Bred by the Cookies Fam. Named after the candy brand for its intense sweetness. THC 18–25%. Dense colorful buds. Rapid-hitting, well-balanced.",
    ),
    "Donny Burger": dict(
        category="flower",
        strain_type="Indica",
        lineage="GMO (Garlic Cookies) x Han-Solo Burger (Skunk House Genetics)",
        effects="Relaxing, Sleepy, Body High, Euphoric, Hungry, Calming, Tingly",
        flavors="Ammonia, Lemon, Pepper, Spice, Sweet, Pungent, Rubber, Cheese",
        terpenes="Myrcene, Pinene, Limonene, Caryophyllene",
        therapeutic="Inflammation, Chronic Pain, Insomnia, Stress, Appetite Stimulation",
        negative="Dry Mouth & Eyes, Heavy Sedation — Not for Daytime. Can be Overwhelming for Novices.",
        aroma="Pungent gas and diesel, spicy savory garlic notes, sour lemon, hints of cheese and fuel",
        misc="Also known as GMO BX. One of the hardest-hitting indicas on the market. THC 25–30%+. Fast-acting, long-lasting body stone. Best reserved for evening. Not for new consumers.",
    ),
    "Super Boof": dict(
        category="flower",
        strain_type="Hybrid",
        lineage="Black Cherry Punch x Tropicana Cookies (Franchise Genetics)",
        effects="Euphoric, Uplifting, Creative, Focused, Happy, Social, Relaxing",
        flavors="Citrus, Earthy, Sweet, Grapefruit, Orange, Sour, Cherry",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Stress, Depression, Fatigue, Mild Pain, Creative Blocks",
        negative="Dry Mouth & Eyes, Occasional Mild Anxiety",
        aroma="Fresh orange and grapefruit citrus burst, sweet black cherry, earthy tropical undertones, faint cookie dough base",
        misc="Bred by Franchise Genetics. One of the most vibrant citrus terpene profiles available. THC 18–24%. Excellent daytime or social strain — uplifting, creative, mood-enhancing.",
    ),
    "Grape Canyon Zkittlez": dict(
        category="flower",
        strain_type="Hybrid (Indica)",
        lineage="Grape Ape x Zkittlez (aka Grape Zkittlez)",
        effects="Euphoric, Relaxing, Happy, Social, Calming, Unbothered",
        flavors="Fruity, Berry, Citrus, Sweet, Sour, Grape",
        terpenes="Myrcene, Linalool, Limonene, Caryophyllene, Humulene",
        therapeutic="Stress, Anxiety, Mild Pain, Mood Enhancement, Appetite",
        negative="Dry Mouth & Eyes, Potential Drowsiness at Higher Doses",
        aroma="Intense grape candy, mixed berry compote, sweet citrus, tropical Zkittlez notes, light earthy base",
        misc="Combines Grape Ape's deep grape flavor with Zkittlez' candy sweetness. THC 19–24%. Dense colorful buds with purple hues. Well-balanced — calming without being fully sedating.",
    ),
    "GG4": dict(
        category="flower",
        strain_type="Hybrid (Indica)",
        lineage="Chem's Sister x Sour Dubb x Chocolate Diesel (GG Strains)",
        effects="Relaxing, Sleepy, Euphoric, Body High, Hungry, Happy",
        flavors="Chemical, Chocolate, Coffee, Diesel, Pine, Sweet",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Chronic Pain, Insomnia, PTSD, Depression, Muscle Spasms",
        negative="Dry Mouth & Eyes, Heavy Sedation — Can be Overwhelming for Novice Users",
        aroma="Pungent diesel and mocha chocolate, earthy pine resin, coffee grounds, chemical fuel — one of the most recognizable aromas in cannabis",
        misc="One of the most iconic strains ever created. Multiple Cannabis Cup winner. Bred by GG Strains. THC 25–32%. Extremely sticky, resin-coated buds. Powerful and fast-acting.",
    ),
    "Grape Cream Cake": dict(
        category="flower",
        strain_type="Hybrid (Indica)",
        lineage="Grape Pie x Wedding Crasher x Ice Cream Cake (Bloom Seed Co.)",
        effects="Relaxing, Sleepy, Euphoric, Happy, Body High, Calming",
        flavors="Fruity, Sweet, Vanilla, Cherry, Nutty, Grape, Creamy",
        terpenes="Myrcene, Linalool, Caryophyllene",
        therapeutic="Chronic Pain, Insomnia, Stress, Anxiety Reduction",
        negative="Couch-Lock at High Doses, Dry Mouth & Eyes",
        aroma="Sweet grape soda, ripe berry, whipped vanilla cream, doughy sweetness with faint gasoline on exhale",
        misc="Bred by Harry Palms (Bloom Seed Co.), former head breeder at Oni Seed Co. THC 20–29%. Dense resinous buds with purple hues and heavy trichome coverage. Best for evening use.",
    ),
    "Gelato 45": dict(
        category="flower",
        strain_type="Hybrid",
        lineage="Girl Scout Cookies x Sherbet (Cookies Fam Genetics — Gelato #45 Phenotype)",
        effects="Euphoric, Calming, Giggly, Happy, Uplifting, Social",
        flavors="Fruity, Berry, Citrus, Pepper, Sweet, Woody, Sour, Creamy",
        terpenes="Limonene, Caryophyllene, Humulene",
        therapeutic="Stress, Depression, Chronic Pain, PTSD, Nausea",
        negative="Dry Mouth & Eyes, Mild Sedation at High Doses",
        aroma="Sweet berry and orange sherbet, creamy vanilla, herbal pepper, faint woody kush undertone",
        misc="One of the most prized phenotypes from the original Gelato line by the Cookies Fam. THC 15–20%. World-class terpene expression and euphoric effects. Balanced hybrid that reads more indica at higher doses.",
    ),

    # ── PRE-ROLL ─────────────────────────────────────────────────────────────
    "Blueberry Muffin": dict(
        category="pre-roll",
        strain_type="Hybrid (Indica)",
        lineage="Purple Panty Dropper x Blueberry (Humboldt Seed Co.)",
        effects="Relaxing, Happy, Euphoric, Sleepy, Uplifting, Body High, Calming",
        flavors="Blueberry, Herbal, Sweet, Vanilla, Butter, Nutty",
        terpenes="Myrcene, Linalool, Pinene, Limonene, Caryophyllene, Humulene, B Pinene",
        therapeutic="Stress, Anxiety, Mild Pain, Insomnia, Depression",
        negative="Dry Mouth & Eyes, Mild Couch-Lock at High Doses",
        aroma="Sweet fresh blueberries, baked muffin dough, light vanilla cream, subtle herbal earthiness",
        misc="Created by Humboldt Seed Co. THC 18–23%. Dense colorful buds with deep purple hues. Beginner-friendly and crowd-pleasing. Available as mini pre-rolls (Lifted North) and singles (Grasslandz).",
    ),
    "Halle Berry": dict(
        category="pre-roll",
        strain_type="Hybrid (Sativa)",
        lineage="Ice Cream Cake x Blockberry (Redwood County Weed Co.)",
        effects="Happy, Uplifting, Euphoric, Focused, Creative, Social, Relaxing, Aroused",
        flavors="Fruity, Berry, Citrus, Earthy, Sweet, Orange, Sour, Creamy",
        terpenes="Myrcene, Linalool, Limonene, Caryophyllene",
        therapeutic="Stress, Anxiety, Depression, Mild Pain, Mood Enhancement",
        negative="Dry Mouth & Eyes",
        aroma="Sweet berry cream, orange citrus, tropical fruit, light earthiness, vanilla cream finish",
        misc="Redwood County Weed Co. cultivar. THC 14–18%. Approachable potency — excellent for daytime or new consumers. Well-balanced, pleasant, mood-lifting.",
    ),
    "Slurricane": dict(
        category="pre-roll",
        strain_type="Hybrid (Indica)",
        lineage="Do-Si-Dos x Purple Punch (In House Genetics)",
        effects="Relaxing, Sleepy, Body High, Hungry, Tingly, Euphoric",
        flavors="Fruity, Berry, Herbal, Spice, Sweet, Tropical, Grape, Candy",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Insomnia, Chronic Pain, Stress, Appetite Stimulation, Nausea",
        negative="Dry Mouth & Eyes, Heavy Sedation — Not for Daytime Use",
        aroma="Tropical grape punch, sweet berry candy, herbal earthiness, faint spice — rich and indulgent",
        misc="Bred by In House Genetics. THC 20–28%. Dense purple flowers. Fast-acting, long-lasting. Best for evening or nighttime use.",
    ),
    "Jelly Breath": dict(
        category="pre-roll",
        strain_type="Hybrid (Indica)",
        lineage="Mendo Breath x Do-Si-Dos (Archive Seed Bank)",
        effects="Relaxing, Sleepy, Euphoric, Body High, Calming, Tingly",
        flavors="Fruity, Berry, Sweet, Grape, Caramel, Vanilla, Earthy",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Chronic Pain, Insomnia, Stress, Muscle Tension, Appetite Stimulation",
        negative="Dry Mouth & Eyes, Heavy Couch-Lock",
        aroma="Sweet grape jelly, caramel, vanilla cream, earthy Mendo funk, hint of gas — like a fruity candy bar with a dank finish",
        misc="Archive Seed Bank genetics. THC 14–20%. Dense frosted flowers. Named for its distinctively sweet jelly-like aroma. Great for pain and sleep.",
    ),
    "La Cegua": dict(
        category="pre-roll",
        strain_type="Hybrid (Sativa)",
        lineage="Black Cherry Punch x Hash Church (Redwood County Weed Co.)",
        effects="Relaxing, Happy, Uplifting, Chill, Euphoric, Tingly",
        flavors="Pine, Skunk, Sweet, Cherry, Earthy, Floral",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Stress Relief, Mood Enhancement, Mild Anxiety, Fatigue",
        negative="Dry Mouth & Eyes",
        aroma="Sweet cherry pine, skunky earthiness, floral undertones, light spice — aromatic and approachable",
        misc="Redwood County Weed Co. cultivar. Named after the Costa Rican mythological spirit La Cegua. THC 15–19%. Balanced, approachable hybrid — a good entry point for newer consumers.",
    ),
    "Early Lemon Berry": dict(
        category="pre-roll",
        strain_type="Hybrid",
        lineage="Las Vegas Lemon Skunk x Member Berry (Elev8 Seeds)",
        effects="Uplifting, Creative, Euphoric, Happy, Relaxing, Body High, Aroused",
        flavors="Fruity, Berry, Citrus, Sweet, Sour, Lemon, Blueberry",
        terpenes="Pinene, Caryophyllene, Humulene",
        therapeutic="Stress, Depression, Mild Chronic Pain, Fatigue, Mood Enhancement",
        negative="Dry Mouth & Eyes, Possible Slight Anxiety",
        aroma="Bright lemon citrus, sweet mixed berries, light skunky earthiness, pine, faint floral — vibrant and fresh",
        misc="Elev8 Seeds genetics. THC 12–16%. One of the most approachable potency levels on the menu. Excellent for daytime use or newer consumers. Great starter pre-roll.",
    ),
    "Purple Milk": dict(
        category="pre-roll",
        strain_type="Hybrid (Sativa)",
        lineage="Grape Gas x Cereal Milk (Compound Genetics)",
        effects="Relaxing, Euphoric, Happy, Calming, Sleepy, Blissful",
        flavors="Grape, Berry, Vanilla, Cream, Sweet, Earthy",
        terpenes="Limonene, Caryophyllene, Humulene",
        therapeutic="Stress, Anxiety, Mild Pain, Insomnia, Mood Disorders",
        negative="Dry Mouth & Eyes, Potential Drowsiness at Higher Doses",
        aroma="Grape soda and creamy milk sweetness, berry compote, light vanilla, earthy and slightly gassy base",
        misc="Compound Genetics cut. THC 28–35%+ in premium batches. Dense violet-tinted flowers packed with resin. Great evening strain with dessert-like flavor.",
    ),
    "Froot by the Foot": dict(
        category="pre-roll",
        strain_type="Hybrid (Indica)",
        lineage="Cotton Candy x Mythic OG (Atlas Seeds)",
        effects="Euphoric, Uplifting, Happy, Creative, Social, Relaxing",
        flavors="Fruity, Spice, Sweet, Candy, Tropical",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Stress Relief, Mood Enhancement, Muscle Relaxation, Mild Pain, Anxiety",
        negative="Dry Mouth & Eyes",
        aroma="Vibrant sweet fruit candy, tropical punch, light gassy undertone — room-filling aromatic presence",
        misc="Bred by Atlas Seeds. THC 20–25%+. Fast flowering, high resin production. Named for its nostalgic candy-like terpene profile. Avión craft pre-roll.",
    ),
    "Zesty Parm": dict(
        category="pre-roll",
        strain_type="Hybrid (Sativa)",
        lineage="Cherry Kush x (Lemon Tree Papaya x GMO) x Fro'Do",
        effects="Focused, Euphoric, Creative, Uplifting, Social, Cerebral",
        flavors="Herbal, Lemon, Pepper, Pine, Skunk, Savory, Citrus",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Chronic Pain, Depression, Mood Swings, Arthritis, Muscle Spasms & Cramps",
        negative="Dry Mouth & Eyes, Mild Headache at High Doses",
        aroma="Earthy spice, fresh woods, peppery zest, dank herbal cheese — like lemon-pepper parmesan with a funky GMO kick",
        misc="THC 18–23%, ~1% CBD. Named for its bold savory-citrus parmesan flavor. Dense olive-green buds with amber undertones. Most unique flavor profile on the menu. Avión craft pre-roll.",
    ),
    "Frodo Sqeeze": dict(
        category="pre-roll",
        strain_type="Hybrid (Sativa)",
        lineage="(Fro'Do x Alaskan Yeti) x LilMissHS (Phylos Bioscience)",
        effects="Relaxing, Happy, Uplifting, Chill, Euphoric, Creative",
        flavors="Fruity, Pine, Spice, Earthy, Herbal",
        terpenes="Myrcene, Limonene, Terpinene",
        therapeutic="Stress Relief, Mood Enhancement, Mild Anxiety, Relaxation",
        negative="Dry Mouth & Eyes",
        aroma="Earthy and fruity with herbal lavender notes, light floral, subtle spice — bright and approachable",
        misc="Phylos Bioscience autoflower genetics. THC ~20–25%. High resin production and stacking cola structure. Balanced hybrid — calm but not sedating. Avión craft pre-roll.",
    ),

    # ── VAPES ────────────────────────────────────────────────────────────────
    "Sour Apple": dict(
        category="vapes",
        strain_type="Hybrid — 76.8% THC",
        lineage="Sour Diesel x Cinderella 99",
        effects="Relaxing, Social, Creative, Body High, Hungry, Happy, Uplifting",
        flavors="Fruity, Citrus, Lemon, Apple, Sweet, Lime, Pear",
        terpenes="Myrcene, Pinene, Caryophyllene",
        therapeutic="Stress Relief, Appetite Stimulation, Mild Depression, Fatigue, Social Anxiety",
        negative="Dry Mouth & Eyes, Mild Sedation at High Doses",
        aroma="Tart green apple candy, lemon-lime citrus, sweet berry, light diesel in background, tropical fruit finish",
        misc="Combines Sour Diesel's cerebral buzz with Cinderella 99's fruity sativa energy. At 76.8% THC. Sweet, clean vape experience — lighter and more approachable than pure Sour Diesel. Balanced, social, and uplifting.",
    ),
    "Sour Diesel": dict(
        category="vapes",
        strain_type="Hybrid (Sativa) — 80.3% THC",
        lineage="Mexican Sativa x Chemdawg (NYC underground, 1990s)",
        effects="Energetic, Euphoric, Creative, Uplifting, Focused, Happy",
        flavors="Citrus, Diesel, Skunk, Sweet, Candy",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Depression, Fatigue, Stress, ADHD, Chronic Pain",
        negative="Dry Mouth & Eyes, Anxiety or Racing Thoughts in Sensitive Users",
        aroma="Sharp diesel fuel, pungent lemon citrus, skunk, earthy pine — one of the most recognizable cannabis aromas of the past 30 years",
        misc="A legendary sativa from 1990s NYC. At 80.3% THC in vape form delivers Sour Diesel's iconic cerebral rush in concentrated form. Fast-acting, energizing, creative. Perfect for daytime. One of the most culturally significant cannabis strains ever produced.",
    ),
    "Girl Scout Cookies": dict(
        category="vapes",
        strain_type="Hybrid — 80.5% THC",
        lineage="OG Kush x Durban Poison (Cookies Fam, San Francisco)",
        effects="Euphoric, Relaxing, Happy, Creative, Focused, Body High",
        flavors="Lemon, Sweet, Woody, Cherry, Mint, Cookie Dough",
        terpenes="Myrcene, Limonene, Caryophyllene",
        therapeutic="Chronic Pain, PTSD, Depression, Nausea, Appetite Stimulation",
        negative="Dry Mouth & Eyes, Possible Paranoia at High Doses",
        aroma="Sweet minty cookie dough, earthy OG musk, lemon zest, cherry notes, hint of pine",
        misc="One of the most culturally impactful strains of the 2010s — parent of Gelato, Sunset Sherbet, MAC, and Runtz. At 80.5% THC. Complex high: euphoric rush followed by full-body relaxation. Ideal for afternoon or evening.",
    ),
    "White Widow": dict(
        category="vapes",
        strain_type="Hybrid — 81% THC",
        lineage="Brazilian Sativa x South Indian Indica (Green House Seeds, Netherlands, 1994)",
        effects="Euphoric, Uplifting, Creative, Social, Alert, Happy",
        flavors="Pine, Spice, Sweet, Woody, Earthy, Floral",
        terpenes="Myrcene, Pinene, Caryophyllene",
        therapeutic="Depression, Stress, Anxiety, Fatigue, Mild Pain",
        negative="Dry Mouth & Eyes, Possible Anxiety or Racing Thoughts at High Doses",
        aroma="Fresh pine resin, earthy spice, faint sweet floral notes, light sandalwood and cedar",
        misc="Dutch cannabis classic and multiple Cannabis Cup winner from 1994. Green House Seeds original genetics. At 81% THC. Balanced hybrid with slight sativa lean. One of the most globally recognized strains in history.",
    ),
}

CATEGORY_ORDER   = ["flower", "pre-roll", "vapes"]
CATEGORY_HEADERS = {
    "flower":   "FLOWER",
    "pre-roll": "PRE-ROLL",
    "vapes":    "VAPES",
}


def add_section_header(doc: Document, title: str):
    """Bold, centered section divider matching the existing doc aesthetic."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{'─' * 20}   {title}   {'─' * 20}")
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()


def add_strain(doc: Document, name: str, data: dict):
    """Add one strain block in exact Strain_Profiles.docx format."""
    # ── Header line: Bold Name (16pt) + tab + regular '- Type' (14pt) ──
    p = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.size = Pt(16)
    r2 = p.add_run("\t")
    r2.bold = True
    r2.font.size = Pt(16)
    r3 = p.add_run(f"-   {data['strain_type']}")
    r3.bold = False
    r3.font.size = Pt(14)

    # ── Field rows: Bold label + regular value ──
    fields = [
        ("Lineage",      data.get("lineage", "")),
        ("Effects",      data.get("effects", "")),
        ("Flavors",      data.get("flavors", "")),
        ("Terpenes",     data.get("terpenes", "")),
        ("Therapeutic",  data.get("therapeutic", "")),
        ("Negative",     data.get("negative", "")),
        ("Aroma",        data.get("aroma", "")),
        ("Misc.",        data.get("misc", "")),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        rb = p.add_run(f"{label}: ")
        rb.bold = True
        rv = p.add_run(value)
        rv.bold = False

    # ── Blank separator line ──
    doc.add_paragraph()


def run():
    doc = Document(SRC)

    # Group new strains by category
    by_cat: dict[str, list] = {c: [] for c in CATEGORY_ORDER}
    for name, data in STRAINS.items():
        if name in ALREADY_IN_DOC:
            print(f"  skip (already in doc): {name}")
            continue
        by_cat[data["category"]].append((name, data))

    for cat in CATEGORY_ORDER:
        items = by_cat[cat]
        if not items:
            continue
        add_section_header(doc, CATEGORY_HEADERS[cat])
        for name, data in items:
            print(f"  adding [{cat}]: {name}")
            add_strain(doc, name, data)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"\nSaved → {OUT}")


if __name__ == "__main__":
    run()
